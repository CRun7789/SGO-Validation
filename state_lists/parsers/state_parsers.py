"""
State-specific parsers for AL (HTML + PDF), AZ, KS, NV (PDF), FL, LA, SC
(HTML), and VA (Word).

PDF states (AZ, KS, NV) pack multiple fields into single text runs that the
generic parser cannot split into columns.  Each uses regex to extract fields
from the concatenated text.

AL lists each SGO on the annual-public-report page as <li><a href="…pdf">Name
</a></li>, where each PDF is that org's annual report filed with the Alabama
DOR.  Section I of every report contains a standardised block with mailing
address, telephone, and email.  The parser fetches all PDFs concurrently;
digitally-filled reports yield full contact info while scanned/image-only
reports (pdfplumber returns no text) degrade gracefully to name-only records.
Phone numbers in digitally-filled reports are rendered with spaced digits
inside the area-code parentheses (PDF form-field artifact) and are
normalised to (NXX) NXX-XXXX format by stripping non-digit characters and
re-formatting the resulting 10-digit string.

FL uses a hand-saved HTML page whose orgs are not in a table or <li> list but
in individual <p> tags, each containing an external <a> link (org name +
website), address/phone as <br/>-separated text, and an obfuscated email via
javascript:mt('user','domain','','').

LA lists orgs as <a class="quickLink"> elements inside a
<div class="QuickLinkList__LinkList"> widget — no table, no <li> items.
Each anchor text is the org name; the href is the org website.

SC has a single SGO (Exceptional SC).  The ECENC program page does not list
it in a table or dedicated element; the org name and website appear in body
prose.  The parser locates the anchor whose href contains "exceptionalsc.org"
and extracts the name and URL from it.

VA provides a Word document with a single table (38 data rows + header):
  Col 0  Name of Approved Scholarship Foundation
  Col 1  Telephone
  Col 2  Web address   (may be "N/A" or have leading non-breaking spaces)
  Col 3  City, State, ZIP Code
  Col 4  Effective Date  (not stored — not in the SGO schema)

The address field is set to the city/state/zip cell (no street address is
provided in the source document).

Fields extracted by state:
  AL — name, address, phone, email    (no EIN, no website; contact info from
                                       individual annual-report PDFs)
  AZ — name, address, phone, website  (no EIN, no email)
  FL — name, address, phone, email, website
  KS — name, address, phone, email    (no EIN, no website)
  LA — name, website                  (no EIN, address, phone, or email)
  NV — name, address, phone, email    (no EIN; website unreliable — omitted)
  SC — name, website                  (no EIN, address, phone, or email)
  VA — name, address, phone, website  (no EIN, no email)
"""
import io
import re
from concurrent.futures import ThreadPoolExecutor, as_completed
from models import SGO

try:
    import pdfplumber
except ImportError as e:
    raise ImportError("pdfplumber is required: pip install pdfplumber") from e

try:
    from bs4 import BeautifulSoup
except ImportError as e:
    raise ImportError("beautifulsoup4 is required: pip install beautifulsoup4") from e

MAX_BYTES = 50_000_000

# ── shared patterns ──────────────────────────────────────────────────────────

_PHONE_RE    = re.compile(r"\(?\d{3}\)?[\s.-]?\d{3}[\s.-][\d\w]{4}")
_EMAIL_RE    = re.compile(r"[\w.+\-]+@[\w\-]+\.[a-z]{2,}", re.IGNORECASE)
# State abbreviation + zip (no word-boundary requirement after digits — KS glues CEO name to zip)
_STATE_ZIP_RE = re.compile(r"([A-Z]{2})\s+(\d{5})")
# First street-number token inside a string (2–6 digits surrounded by spaces)
_STREET_START_RE = re.compile(r"\s+\d{2,6}\s+")

HEADER_WORDS = {
    "name", "sto", "sgo", "mailing", "phone", "website", "dates", "address",
    "city", "state", "zip", "ceo", "contact", "email", "telephone", "certified",
    "tax", "organizations",
}

_AL_MAX_PDF_WORKERS = 10  # concurrent PDF fetches for AL annual reports


def _parse_al_section_i(pdf_bytes: bytes) -> tuple[str | None, str | None, str | None]:
    """
    Extract (address, phone, email) from Section I of an AL SGO annual report.

    Alabama DOR uses a standardised one-page form.  Digitally-filled PDFs have
    an embedded text layer that pdfplumber can read directly; scanned/paper
    submissions have no text layer and return empty strings.  In both cases we
    only inspect page 1 (Section I is always on the first page).

    Phone normalisation: digitally-filled forms render each digit of the area
    code with a trailing space (e.g. "( 2 0 5 ) 2067804").  We strip all
    non-digit characters and reformat the 10-digit result as (NXX) NXX-XXXX.
    Already-formatted strings (e.g. "(205) 445-2908") survive the same path
    correctly since their digit sequence is also 10 digits.

    Returns (None, None, None) when no text is extractable (scanned PDF).
    """
    try:
        with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
            if not pdf.pages:
                return None, None, None
            text = pdf.pages[0].extract_text() or ""
    except Exception:
        return None, None, None

    if not text.strip():
        return None, None, None  # scanned image PDF — graceful degradation

    address: str | None = None
    phone:   str | None = None
    email:   str | None = None

    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            continue

        # ── Email + Phone ────────────────────────────────────────────────────
        # Section I prints both on the same line: "<phone digits>  <email>"
        # Skip the column-header row ("TELEPHONE NUMBER   EMAIL ADDRESS").
        if "@" in line and email is None and "EMAIL" not in line.upper()[:20]:
            email_m = _EMAIL_RE.search(line)
            if email_m:
                email = email_m.group()
                # Everything before the email match is the phone field.
                digits = re.sub(r"\D", "", line[:email_m.start()])
                if len(digits) >= 10:
                    d = digits[-10:]  # last 10 in case of stray leading digits
                    phone = f"({d[0:3]}) {d[3:6]}-{d[6:10]}"

        # ── Mailing Address ───────────────────────────────────────────────────
        # The data row contains a 5-digit ZIP code; the label row does not.
        # We also require at least one lowercase letter so we don't accidentally
        # match all-caps header rows that happen to contain a state abbreviation.
        if address is None and re.search(r"\d{5}", line) and re.search(r"[a-z]", line):
            zip_m = re.search(r"\b\d{5}\b", line)
            if zip_m:
                # Keep only the text up through the ZIP code (ignore anything
                # after, such as continuation of a PDF column layout).
                candidate = " ".join(line[:zip_m.end()].split())
                # Require at least one non-digit character before the ZIP so
                # we don't capture stray numeric lines.
                if re.search(r"[A-Za-z]", candidate):
                    address = candidate

    return address, phone, email


_AL_YEAR_RE = re.compile(r"/(\d{4})/")  # matches the year folder in AL DOR upload URLs


def parse_html_al(html: str, state: str, url: str) -> list[SGO]:
    """
    Extract SGOs from the Alabama DOR annual-public-report page.

    Strategy:
      1. Parse the HTML for every <a href="…pdf"> anchor whose href points to
         an AL DOR annual SGO report.  Extract the report year from the URL's
         upload-folder path segment (e.g. /uploads/2025/10/…).
      2. Determine the most recent year present across all collected links and
         keep only those links.  This excludes orgs that filed in prior years
         but have not yet filed for the current year.
      3. Fetch all current-year report PDFs concurrently (up to
         _AL_MAX_PDF_WORKERS threads).
      4. Parse Section I of each PDF for mailing address, phone, and email.
         Scanned / image-only PDFs return (None, None, None) and are stored
         as name-only records without raising an error.

    The `fetch_bytes` function is imported lazily from `fetcher` to avoid a
    circular import at module load time (state_parsers ← fetcher ← sources).
    """
    soup = BeautifulSoup(html, "html.parser")

    # Remove nav chrome so we only see content anchors
    for tag in soup.find_all(["nav", "header", "footer"]):
        tag.decompose()

    # Collect (name, pdf_url, year) for every AL DOR annual SGO report link
    candidates: list[tuple[str, str, int]] = []

    for a in soup.find_all("a", href=True):
        href = a["href"].strip()
        if not href.lower().endswith(".pdf"):
            continue
        if "revenue.alabama.gov" not in href and not href.startswith("/"):
            continue
        name = a.get_text(strip=True)
        if not name or len(name) < 3:
            continue
        year_m = _AL_YEAR_RE.search(href)
        year = int(year_m.group(1)) if year_m else 0
        candidates.append((name, href, year))

    if not candidates:
        raise ValueError(f"No SGO PDF links found on AL page: {url}")

    # Keep only links from the most recent year
    most_recent_year = max(year for _, _, year in candidates)
    print(f"  [AL] Most recent report year: {most_recent_year} "
          f"({sum(1 for _, _, y in candidates if y == most_recent_year)} orgs)")

    # Deduplicate by normalised name within the most-recent-year set
    # (guard against duplicate anchors for the same org on the same page)
    seen_names: set[str] = set()
    org_links: list[tuple[str, str]] = []
    for name, href, year in candidates:
        if year != most_recent_year:
            continue
        key = re.sub(r"\W+", "", name).lower()
        if key in seen_names:
            continue
        seen_names.add(key)
        org_links.append((name, href))

    if not org_links:
        raise ValueError(f"No SGO PDF links found on AL page: {url}")

    # Lazy import to avoid circular dependency at module load time
    try:
        from fetcher import fetch_bytes as _fetch_bytes
    except ImportError:
        _fetch_bytes = None  # type: ignore[assignment]

    def _fetch_and_parse(item: tuple[str, str]) -> SGO:
        org_name, pdf_url = item
        addr = ph = em = None
        if _fetch_bytes is not None:
            try:
                pdf_bytes = _fetch_bytes(pdf_url)
                addr, ph, em = _parse_al_section_i(pdf_bytes)
            except Exception:
                pass  # network error or corrupt PDF — return name-only record
        return SGO(
            state=state,
            name=org_name,
            ein=None,
            raw_source=url,
            address=addr,
            phone=ph,
            email=em,
            website=None,
        )

    results: list[SGO] = []
    with ThreadPoolExecutor(max_workers=_AL_MAX_PDF_WORKERS) as pool:
        futures = {pool.submit(_fetch_and_parse, item): item for item in org_links}
        for future in as_completed(futures):
            try:
                results.append(future.result())
            except Exception:
                org_name, pdf_url = futures[future]
                results.append(SGO(state=state, name=org_name, ein=None, raw_source=url))

    if not results:
        raise ValueError(f"No SGO names extracted from AL page: {url}")

    return results


def _is_header(text: str) -> bool:
    words = text.lower().split()
    return bool(words) and words[0] in HEADER_WORDS


def _split_name_address(text: str) -> tuple[str, str | None]:
    """Return (name, street_address) by splitting at the first street number."""
    m = _STREET_START_RE.search(text)
    if not m:
        return text.strip(), None
    name = text[: m.start()].strip()
    rest = text[m.start():]
    zip_m = _STATE_ZIP_RE.search(rest)
    address = rest[: zip_m.end()].strip() if zip_m else None
    return name, address


# ── AZ ───────────────────────────────────────────────────────────────────────
#
# Each PDF table row is collapsed by pdfplumber to a single string:
#   'OrgName StreetAddr Phone Website DateRange [NameContinued] City, AZ Zip'
#
# Strategy: strip phone, website, and date-range tokens; split on the first
# street number to get name and address; rejoin the city+state+zip from the
# end of the stripped string.

_AZ_DATE_RE    = re.compile(r"\d{2}/\d{2}/\d{2,4}\s*[–—-]\s*(?:Present|Past|\d{2}/\d{2}/\d{2,4})", re.IGNORECASE)
_AZ_PHONE_RE   = re.compile(r"\d{3}-\d{3}-\d{4}")
_AZ_WEBSITE_RE = re.compile(r"\b[\w-]+\.[a-z]{2,}\b", re.IGNORECASE)


def parse_pdf_az(pdf_bytes: bytes, state: str, url: str) -> list[SGO]:
    if len(pdf_bytes) > MAX_BYTES:
        raise ValueError(f"PDF from {url} exceeds size limit")

    results: list[SGO] = []

    with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
        for page in pdf.pages:
            for table in page.extract_tables():
                for row in table:
                    cell = next((c for c in row if c and c.strip()), None)
                    if not cell:
                        continue
                    text = " ".join(cell.split())
                    if _is_header(text):
                        continue

                    # Extract fields before stripping them from the text
                    phone_m = _AZ_PHONE_RE.search(text)
                    phone: str | None = phone_m.group() if phone_m else None

                    web_m = _AZ_WEBSITE_RE.search(text)
                    website: str | None = web_m.group() if web_m else None

                    # Remove phone, website, and date noise to isolate name+address
                    stripped = text
                    if phone_m:
                        stripped = stripped.replace(phone_m.group(), "")
                    if web_m:
                        stripped = stripped.replace(web_m.group(), "")
                    stripped = _AZ_DATE_RE.sub("", stripped)
                    stripped = " ".join(stripped.split())

                    name, address = _split_name_address(stripped)
                    if not name or _is_header(name):
                        continue

                    # If there's still a city+state+zip at the very end that wasn't
                    # captured in address (it occurs after the date-range fragment),
                    # append it to the address.
                    if address is None:
                        city_m = _STATE_ZIP_RE.search(stripped)
                        if city_m:
                            space_pos = stripped.rfind(" ", 0, city_m.start())
                            addr_start = space_pos + 1 if space_pos >= 0 else 0
                            address = stripped[addr_start: city_m.end()].strip()

                    # Remove stray corporate-suffix tokens (Inc., LLC, etc.) that PDF
                    # column reflow inserts between street and city.
                    if address:
                        address = re.sub(r"\s+(?:Inc|LLC|Ltd|Corp|LLP)\.?\s+", " ", address).strip()

                    try:
                        results.append(SGO(
                            state=state, name=name, ein=None, raw_source=url,
                            address=address or None,
                            phone=phone,
                            email=None,
                            website=website,
                        ))
                    except ValueError:
                        pass

    if not results:
        raise ValueError(f"No SGO names extracted from AZ PDF: {url}")

    return results


# ── KS ───────────────────────────────────────────────────────────────────────
#
# Each PDF row is a single string:
#   'OrgName StreetAddr City ST ZipCEOName ContactName Phone Email'
#
# The zip is glued directly to the CEO first name (no space), so we cannot
# rely on a word boundary after the 5-digit zip.

def parse_pdf_ks(pdf_bytes: bytes, state: str, url: str) -> list[SGO]:
    if len(pdf_bytes) > MAX_BYTES:
        raise ValueError(f"PDF from {url} exceeds size limit")

    results: list[SGO] = []

    # KS PDF needs the Referer header — use a fresh session-like call via pdfplumber
    with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
        for page in pdf.pages:
            for table in page.extract_tables():
                for row in table:
                    cell = next((c for c in row if c and c.strip()), None)
                    if not cell:
                        continue
                    text = " ".join(cell.split())
                    if _is_header(text):
                        continue

                    # Email is the last token (always at the end)
                    email: str | None = None
                    email_m = _EMAIL_RE.search(text)
                    if email_m:
                        email = email_m.group()
                        text_no_email = text[: email_m.start()].rstrip()
                    else:
                        text_no_email = text

                    # Phone: last match before the email
                    phone: str | None = None
                    phone_ms = list(_PHONE_RE.finditer(text_no_email))
                    if phone_ms:
                        phone = phone_ms[-1].group()

                    # Name: everything up to the first street number
                    street_m = _STREET_START_RE.search(text)
                    if street_m:
                        name = text[: street_m.start()].strip()
                        addr_start = street_m.start()
                        # Address: from street number through state+zip
                        # _STATE_ZIP_RE has no \b after digits, so it matches "80111Norton"
                        zip_m = _STATE_ZIP_RE.search(text, addr_start)
                        if zip_m:
                            address: str | None = text[addr_start: zip_m.end()].strip()
                        else:
                            address = None
                    else:
                        name = text_no_email.strip()
                        address = None

                    if not name or _is_header(name):
                        continue

                    try:
                        results.append(SGO(
                            state=state, name=name, ein=None, raw_source=url,
                            address=address,
                            phone=phone,
                            email=email,
                            website=None,
                        ))
                    except ValueError:
                        pass

    if not results:
        raise ValueError(f"No SGO names extracted from KS PDF: {url}")

    return results


# ── NV ───────────────────────────────────────────────────────────────────────
#
# Free-text, multi-line blocks.  Each org occupies several consecutive lines:
#   Line 1:        OrgName  (NNN)NNN-XXXX
#   Lines 2…N:     Contact Person, Title  contact@email.com   [+ fax lines]
#   Near-last:     Street address   [fax]
#   Last (maybe):  Website text
#
# Block boundary: a line that (a) has a parenthesized phone, (b) does NOT end
# with "(fax)", and (c) does NOT contain a comma (which marks "Person, Title").
# The one org without a phone on its name line (IPOF) is handled by treating
# consecutive non-address, non-contact lines as name candidates.

_NV_PHONE_PAREN_RE = re.compile(r"\(\d{3}\)\d{3}[-\s][\d\w]{4}")
_NV_FAX_RE = re.compile(r"\(fax\)\s*$", re.IGNORECASE)


def _is_nv_org_line(line: str) -> bool:
    """True if this line is an org-name line (phone present, not a fax/contact line).

    NV contact-person lines always end with an email address; org-name lines
    never do.  Using email absence is more robust than a fixed title-word list
    (which would miss titles like 'VP', 'Superintendent', 'Liaison', etc.).
    """
    return (
        bool(_NV_PHONE_PAREN_RE.search(line))
        and not _NV_FAX_RE.search(line)
        and not _EMAIL_RE.search(line)
    )


def parse_pdf_nv(pdf_bytes: bytes, state: str, url: str) -> list[SGO]:
    if len(pdf_bytes) > MAX_BYTES:
        raise ValueError(f"PDF from {url} exceeds size limit")

    all_lines: list[str] = []
    with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            all_lines.extend(ln.strip() for ln in text.splitlines() if ln.strip())

    # Split into blocks on org-name lines
    blocks: list[list[str]] = []
    current: list[str] = []
    for line in all_lines:
        if _is_nv_org_line(line) and current:
            blocks.append(current)
            current = [line]
        else:
            current.append(line)
    if current:
        blocks.append(current)

    results: list[SGO] = []
    for block in blocks:
        if not block:
            continue

        first = block[0]

        # Skip page headers / boilerplate that slipped into the first block
        if _is_header(first) or not first or len(first) < 4:
            continue

        # Phone: from the org-name line
        phone_m = _NV_PHONE_PAREN_RE.search(first)
        phone: str | None = phone_m.group() if phone_m else None
        name = first[: phone_m.start()].strip() if phone_m else first.strip()

        if not name or _is_header(name):
            continue

        # Address: first line in the block that starts with a digit
        address: str | None = None
        for line in block[1:]:
            if re.match(r"^\d", line):
                # Strip trailing fax number from the address line
                addr = re.sub(r"\s+\(?\d{3}\)?[\s.-]?\d{3}[\s.-][\d\w]{4}.*$", "", line).strip()
                # Also strip any repeated org-name text appended after city/zip
                zip_m = _STATE_ZIP_RE.search(addr)
                if zip_m:
                    addr = addr[: zip_m.end()].strip()
                address = addr or None
                break

        # Email: first email in any line after the org-name line
        email: str | None = None
        for line in block[1:]:
            em = _EMAIL_RE.search(line)
            if em:
                email = em.group()
                break

        try:
            results.append(SGO(
                state=state, name=name, ein=None, raw_source=url,
                address=address,
                phone=phone,
                email=email,
                website=None,
            ))
        except ValueError:
            pass

    if not results:
        raise ValueError(f"No SGO names extracted from NV PDF: {url}")

    return results


# ── FL ───────────────────────────────────────────────────────────────────────
#
# The FL SFO page lists each org as a standalone <p> block:
#
#   <p>
#     <a href="https://orgwebsite.org/" target="_blank">Org Name</a><br/>
#     P.O. Box 123, City, FL 00000<br/>
#     Phone/Fax: 555-555-5555<br/>
#     <a href="javascript:mt('user','domain.org','','')">user@domain.org</a>
#   </p>
#
# The email address is obfuscated via a JS function call; we reconstruct it
# from the mt() arguments: mt('localpart', 'domain', ...) → localpart@domain.
# Some orgs omit the email or phone; all have a name link.

_FL_MT_RE = re.compile(r"javascript:mt\('([^']+)','([^']+)'", re.IGNORECASE)
# Labeled: "Phone/Fax: 888-707-2465"
_FL_PHONE_LABELED_RE = re.compile(r"(?:Phone|Fax|Tel|Phone/Fax)[^:]*:\s*([\d()\s.\-/]+)", re.IGNORECASE)
# Bare: a line that is nothing but digits, spaces, dashes, parens — at least 7 chars
_FL_PHONE_BARE_RE = re.compile(r"^[\d()\s.\-/]{7,}$")


def parse_html_fl(html: str, state: str, url: str) -> list[SGO]:
    """
    Extract SGOs from the FL SFO page (manually downloaded HTML).

    Finds the content div by id='newPageContent', then iterates every <p>
    that contains an external <a> link as its first anchor.
    """
    soup = BeautifulSoup(html, "html.parser")

    # Prefer the specific content div; fall back to the article, then body
    root = (
        soup.find("div", id="newPageContent")
        or soup.find("article", id="inner-content")
        or soup.body
    )
    if root is None:
        raise ValueError(f"Could not find content area in FL HTML: {url}")

    results: list[SGO] = []

    for p in root.find_all("p"):
        anchors = p.find_all("a", href=True)
        if not anchors:
            continue

        # First anchor with an external href is the org name + website
        name_tag = next(
            (a for a in anchors if a["href"].startswith(("http://", "https://"))),
            None,
        )
        if name_tag is None:
            continue

        name = name_tag.get_text(strip=True)
        website: str | None = name_tag["href"]

        # Skip paragraphs whose body text is boilerplate rather than contact info
        # (e.g. "Please note: If you are the owner or operator of a private school…")
        p_body = p.get_text(separator=" ", strip=True)
        if p_body.lower().startswith("please note"):
            continue

        # Email: reconstruct from javascript:mt('user','domain','','')
        email: str | None = None
        for a in anchors:
            m = _FL_MT_RE.match(a.get("href", ""))
            if m:
                email = f"{m.group(1)}@{m.group(2)}"
                break

        # Address + phone: <br/>-separated text segments after the name anchor
        # Collect all NavigableString children (between <br/> tags)
        raw_lines = [
            seg.strip()
            for seg in p.strings
            if seg.strip() and seg.strip() != name
        ]

        phone: str | None = None
        addr_parts: list[str] = []
        for line in raw_lines:
            m = _FL_PHONE_LABELED_RE.match(line)
            if m:
                phone = m.group(1).strip().rstrip("/")
            elif _FL_PHONE_BARE_RE.match(line):
                phone = line.strip()
            elif "@" not in line:           # skip email display text
                addr_parts.append(line)

        address: str | None = ", ".join(addr_parts) if addr_parts else None

        try:
            results.append(SGO(
                state=state, name=name, ein=None, raw_source=url,
                address=address,
                phone=phone,
                email=email,
                website=website,
            ))
        except ValueError:
            pass

    if not results:
        raise ValueError(f"No SGO names extracted from FL HTML: {url}")

    return results


# ── LA ───────────────────────────────────────────────────────────────────────
#
# The LA Tuition Donation Credit page lists each participating STO as an
# <a class="quickLink"> element inside a single widget div:
#
#   <div class="QuickLinkList__LinkList">
#     <a class="quickLink" href="https://orgwebsite.org/">Org Name</a>
#     ...
#   </div>
#
# The anchor text is the org name; the href is the org website.
# No address, phone, or email is present on the page.

def parse_html_la(html: str, state: str, url: str) -> list[SGO]:
    """
    Extract STOs from the LA Tuition Donation Credit program page.

    Targets the QuickLinkList widget that holds each participating org as an
    <a class="quickLink"> anchor; falls back to any quickLink anchor on the page
    if the wrapper div is absent.
    """
    soup = BeautifulSoup(html, "html.parser")

    link_list = soup.find("div", class_="QuickLinkList__LinkList")
    if link_list:
        anchors = link_list.find_all("a", class_="quickLink")
    else:
        # Fallback: any quickLink anchor on the page
        anchors = soup.find_all("a", class_="quickLink")

    results: list[SGO] = []
    for a in anchors:
        name = a.get_text(strip=True)
        if not name:
            continue
        website: str | None = a.get("href") or None
        try:
            results.append(SGO(
                state=state, name=name, ein=None, raw_source=url,
                address=None, phone=None, email=None, website=website,
            ))
        except ValueError:
            pass

    if not results:
        raise ValueError(f"No SGO names extracted from LA HTML: {url}")

    return results


# ── SC ───────────────────────────────────────────────────────────────────────
#
# The SC ECENC program page does not have a dedicated org list.  Exceptional SC
# is mentioned in body prose, and its website (exceptionalsc.org) appears as an
# <a> anchor.  We locate that anchor to get the org name and URL rather than
# hardcoding them, so the parser will still work if the page wording changes.

def parse_html_sc(html: str, state: str, url: str) -> list[SGO]:
    """
    Extract the single SGO from the SC ECENC program page (manually downloaded).

    Finds the <a> anchor whose href points to exceptionalsc.org and uses its
    link text as the org name.  Falls back to a hardcoded name if the anchor
    is not found (e.g. the page was restructured).
    """
    soup = BeautifulSoup(html, "html.parser")

    anchor = soup.find("a", href=lambda h: h and "exceptionalsc.org" in h)
    if anchor:
        name = anchor.get_text(strip=True) or "Exceptional SC"
        website: str | None = anchor["href"]
    else:
        # Page restructured — fall back to the known name with no website URL
        name = "Exceptional SC"
        website = None

    return [
        SGO(
            state=state,
            name=name,
            ein=None,
            raw_source=url,
            address=None,
            phone=None,
            email=None,
            website=website,
        )
    ]


# ── VA ───────────────────────────────────────────────────────────────────────
#
# The VA VDOE Word document contains a single table with 5 columns:
#   0  Name of Approved Scholarship Foundation
#   1  Telephone
#   2  Web address      (may be "N/A" or have leading \xa0 non-breaking spaces)
#   3  City, State, ZIP Code
#   4  Effective Date   (informational only — not stored in the SGO schema)
#
# There is no street address in the source document; the address field is
# populated from the City/State/ZIP cell.
#
# Edge cases handled:
#   - "N/A" website → None
#   - Leading \xa0 before a URL → stripped
#   - Extra whitespace / non-breaking spaces in city/state/zip → normalized

try:
    import docx as _docx
except ImportError as e:
    raise ImportError("python-docx is required: pip install python-docx") from e

MAX_BYTES_DOCX = 50_000_000

_VA_HEADER_FIRST_WORDS = {"name", "scholarship", "organization", "approved"}


def _va_cell(row, col_idx: int) -> str:
    """Return stripped, \xa0-cleaned text from a table cell."""
    if col_idx >= len(row.cells):
        return ""
    return row.cells[col_idx].text.replace("\xa0", " ").strip()


def parse_docx_va(docx_bytes: bytes, state: str, url: str) -> list[SGO]:
    """
    Extract SGOs from the VA VDOE Word document.

    Reads the first table in the document; each data row provides name,
    phone, website, and city/state/zip address.
    """
    if len(docx_bytes) > MAX_BYTES_DOCX:
        raise ValueError(f"docx from {url} exceeds size limit")

    doc = _docx.Document(io.BytesIO(docx_bytes))

    # Find the first table (there is only one)
    if not doc.tables:
        raise ValueError(f"No tables found in VA docx: {url}")

    table = doc.tables[0]
    results: list[SGO] = []

    for row in table.rows:
        name = _va_cell(row, 0)
        if not name:
            continue
        # Skip the header row (first word is a known header term)
        first_word = name.split()[0].lower().rstrip(":")
        if first_word in _VA_HEADER_FIRST_WORDS:
            continue

        phone_raw = _va_cell(row, 1)
        phone: str | None = phone_raw if phone_raw else None

        website_raw = _va_cell(row, 2)
        # Treat "N/A", blank, or whitespace-only as absent
        if website_raw and website_raw.upper() != "N/A":
            # Ensure a usable URL scheme
            website: str | None = website_raw if "." in website_raw else None
        else:
            website = None

        city_state_zip = _va_cell(row, 3)
        address: str | None = city_state_zip if city_state_zip else None

        try:
            results.append(SGO(
                state=state,
                name=name,
                ein=None,
                raw_source=url,
                address=address,
                phone=phone,
                email=None,
                website=website,
            ))
        except ValueError:
            pass

    if not results:
        raise ValueError(f"No SGO names extracted from VA docx: {url}")

    return results
