"""
Extract SGO names from Excel (.xlsx), CSV, and Word (.docx) files.

Install: pip install openpyxl python-docx
csv and io are stdlib — no install needed.
"""
import csv
import io
from models import SGO

try:
    import openpyxl
except ImportError as e:
    raise ImportError("openpyxl is required: pip install openpyxl") from e

try:
    import docx
except ImportError as e:
    raise ImportError("python-docx is required: pip install python-docx") from e

MAX_BYTES = 50_000_000
HEADER_WORDS = {"name", "organization", "sgo", "scholarship", "no.", "#", "entity", "document"}


def _is_header(text: str) -> bool:
    words = text.lower().split()
    return bool(words) and words[0] in HEADER_WORDS


def parse_xlsx(xlsx_bytes: bytes, state: str, url: str) -> list[SGO]:
    """
    Read the first sheet of an .xlsx file and extract one SGO per row.

    Assumes the first column contains org names; skips the header row.
    """
    if len(xlsx_bytes) > MAX_BYTES:
        raise ValueError(f"xlsx from {url} exceeds size limit")

    wb = openpyxl.load_workbook(io.BytesIO(xlsx_bytes), read_only=True, data_only=True)
    ws = wb.active
    results = []

    for row in ws.iter_rows(values_only=True):
        if not row:
            continue
        first = str(row[0]).strip() if row[0] is not None else ""
        if not first or _is_header(first):
            continue
        try:
            results.append(SGO(state=state, name=first, ein=None, raw_source=url))
        except ValueError:
            pass

    wb.close()

    if not results:
        raise ValueError(f"No SGO names extracted from xlsx: {url}")

    return results


_CSV_COL_ALIASES: dict[str, list[str]] = {
    "name":    ["organization name", "name", "sgo name", "sto name", "organization"],
    "ein":     ["ein", "federal tax id", "tax id", "fein"],
    "address": ["address #1", "address", "mailing address", "street address"],
    "address2":["address #2"],
    "address3":["address #3"],
    "phone":   ["telephone", "phone", "phone number"],
    "email":   ["e-mail", "email"],
    "website": ["web-site", "website", "url", "web address", "sto website"],
}


def _col_index(headers: list[str], aliases: list[str]) -> int | None:
    """Return the index of the first header that matches any alias (case-insensitive)."""
    hl = [h.strip().lower() for h in headers]
    for alias in aliases:
        if alias in hl:
            return hl.index(alias)
    return None


def parse_csv(csv_bytes: bytes, state: str, url: str) -> list[SGO]:
    """
    Read a CSV file and extract one SGO per row.

    Reads the header row to locate name, EIN, address, phone, email, and
    website columns when present; falls back to first-column-only if the
    header row is absent or unrecognised.
    """
    if len(csv_bytes) > MAX_BYTES:
        raise ValueError(f"CSV from {url} exceeds size limit")

    text = csv_bytes.decode("utf-8", errors="replace")
    reader = csv.reader(io.StringIO(text))
    rows = list(reader)
    if not rows:
        raise ValueError(f"No SGO names extracted from CSV: {url}")

    # Detect header row
    col: dict[str, int | None] = {}
    data_start = 0
    first_vals = [v.strip() for v in rows[0]]
    if first_vals and _is_header(first_vals[0]):
        for field, aliases in _CSV_COL_ALIASES.items():
            col[field] = _col_index(first_vals, aliases)
        data_start = 1
    else:
        col = {f: None for f in _CSV_COL_ALIASES}

    def _cell(row: list[str], field: str) -> str | None:
        idx = col.get(field)
        if idx is not None and idx < len(row):
            v = row[idx].strip()
            return v or None
        return None

    results = []
    for row in rows[data_start:]:
        if not row:
            continue
        name = _cell(row, "name") or (row[0].strip() if row else "")
        if not name:
            continue

        # Combine multi-part address fields
        addr_parts = [
            _cell(row, "address"),
            _cell(row, "address2"),
            _cell(row, "address3"),
        ]
        address = ", ".join(p for p in addr_parts if p) or None

        try:
            results.append(SGO(
                state=state,
                name=name,
                ein=_cell(row, "ein"),
                raw_source=url,
                address=address,
                phone=_cell(row, "phone"),
                email=_cell(row, "email"),
                website=_cell(row, "website"),
            ))
        except ValueError:
            pass

    if not results:
        raise ValueError(f"No SGO names extracted from CSV: {url}")

    return results


def parse_docx(docx_bytes: bytes, state: str, url: str) -> list[SGO]:
    """
    Extract SGO names from a Word document.

    Checks tables first (one name per row), then falls back to paragraphs.
    """
    if len(docx_bytes) > MAX_BYTES:
        raise ValueError(f"docx from {url} exceeds size limit")

    doc = docx.Document(io.BytesIO(docx_bytes))
    results = []

    # Tables first
    for table in doc.tables:
        for row in table.rows:
            first = row.cells[0].text.strip() if row.cells else ""
            if not first or _is_header(first):
                continue
            try:
                results.append(SGO(state=state, name=first, ein=None, raw_source=url))
            except ValueError:
                pass

    # Fallback: paragraphs
    if not results:
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text or text.lower() in HEADER_WORDS:
                continue
            try:
                results.append(SGO(state=state, name=text, ein=None, raw_source=url))
            except ValueError:
                pass

    if not results:
        raise ValueError(f"No SGO names extracted from docx: {url}")

    return results
